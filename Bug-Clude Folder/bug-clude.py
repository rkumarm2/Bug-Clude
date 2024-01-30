#pip install playwright
#pip install python-docx
#pip install pytest-playwright

import aide
import tk
from tkinter import ttk
from tkinter import *
import tkinter as tk

import random
from PIL import Image, ImageTk
from tkinter.filedialog import askopenfiles

import time
from datetime import datetime

import docx
from docx import Document
from docx.text.run import *
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from playwright.sync_api import sync_playwright

#=====================================

def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run ()
    r._r.append (hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink

#======================================================================================================

def start_web_scrapping( ip_bugs_list , doc , page , l , xl , user_pid):
    
    aide.submit_statistics(
        pid = user_pid,
        tool_id="1234567890",
        metadata={
            "potential_savings": 1,  # Hours
            "report_savings": True,
        },
    )


    included_bugs_list =[]
    excluded_bugs_list =[]
    only_exc_checker ,  ex_if_sr_count , ex_if_yrs_old , ex_int_srcount , ex_intwithSEVimpact , ex_sev_is , ex_cl_j_h = l[0] , l[1] , l[2] , l[3] , l[4] , l[5] , l[6] 
    #c3               c4           c5             c6         c9                   c12              c13         c14                                                                                          
    s5,s6,s9,s12,s13,s14a,s14b = xl[0] , xl[1] , xl[2] , xl[3] , xl[4] , xl[5] , xl[6] 
    
    now = datetime.now()
    datevar = now.strftime("%D")
    #print(datevar)
    curm , curd, cury = datevar.split("/")
    cur_year = int("20" + cury)


    for var1 in ip_bugs_list:

        link_str = "https://cdetsng.cisco.com/summary/#/defect/"+var1
        page.goto(link_str)
        page.set_default_timeout(120000)
        title_txt = page.locator("//div[@id='title'][@class='h5']").inner_text()
                
        while(title_txt[:10]!=var1):
            #print(title_txt[:3])
            title_txt = page.locator("//div[@id='title'][@class='h5']").inner_text()


        #print("TITLE TEXT: ",title_txt)
        mn= page.locator("//div[@id='expand'][@class='form-group form-group--inline']") 
        title_list = title_txt.split("\n")
        bug_id = title_list[0]
        bug_title = title_list[1]
        other_det_list = title_list[2].split(" ")
        
        in_fnd_ex_fnd = other_det_list[0]
        bug_sev_name = other_det_list[2]
        bug_sev_val = int(other_det_list[3].strip("(Sev").strip(")") )
        bug_status = other_det_list[4]
        bug_vis = " ".join(other_det_list[5:])
        #print("in_fnd_ex_fnd value : ",in_fnd_ex_fnd)
        #print()
        enc_only = page.locator("//div[@id='enclosures'][@class='col-sm-12 text-lg-left']")
        text_of_enc_only = enc_only.inner_text()
        text_of_enc_only_len = len(text_of_enc_only)
        list_of_enc_only = text_of_enc_only.split("\n")
        list_of_enc_only_len = len(list_of_enc_only)
        #print(list_of_enc_only[0])
        #print(list_of_enc_only[1])
        for fff in range(4):
            del list_of_enc_only[2]
        t = list_of_enc_only[2] 
        topic_count = int(t[t.find("(")+1 : t.find(")")])
        
        #Taking data from histogram
        hist = page.locator("//app-histogram")
        #hist = driver.find_element(By.TAG_NAME,"app-histogram")
        hist_text=hist.inner_text() 
        hist_text=hist_text.split("\n")
        bug_init_date = hist_text[4]
        bug_init_date , bug_inti_month , bug_init_year = bug_init_date.split("/")
        bug_last_edit_date = hist_text[-1]
        bug_last_stat_by_hist = hist_text[-4]


        #Taking data from what area
        what = page.locator("//app-what[@id='whatArea']")
        what_text = what.inner_text().split("\n")
        bug_product = what_text[2]
        bug_component = what_text[3]
        bug_feature = what_text[7]


        #Taking data from who area
        who = page.locator("//app-who[@id='whoArea']")
        who_text = who.inner_text().split("\n")
        if "DE-manager:" in who_text:
            bug_de_man = who_text[who_text.index("DE-manager:") +1] 
        else:
            bug_de_man = "Not available"
        if "Engineer:" in who_text:
            bug_eng_man = who_text[who_text.index("Engineer:") +1] 
        else:
            bug_eng_man = "Not available"


        #taking data from version
        ver = page.locator("//app-version[@id='versionArea']")
        ver_text = ver.inner_text().split("\n")

        bug_1st_fnd_ver = ver_text[2]
        bug_int_rel = "Not available"
        for i in ver_text[3:]:
            if "Integrated-releases:" in i:
                bug_int_rel = i 
                continue 
                
                
        # #Fetching service requests count
        sr_ = page.locator("//app-service-requests")
        sr_txt = sr_.inner_text().split("\n")
#         for i in sr_txt:
#             print(i)
        sr_count = int(sr_txt[0].strip("Service Request(").strip(")"))
        #print("SR-count is : ", sr_count)
#         p = doc.add_paragraph(" ")
#         add_hyperlink(p , var1 , link_str)
#         doc.add_paragraph(bug_title)
        ###
        
        if ( ex_if_sr_count==1 and sr_count <= s5 ):
            p = doc.add_paragraph(" ")
            add_hyperlink(p , var1 , link_str)
            doc.add_paragraph(bug_title)
            doc.add_paragraph()
            doc.add_paragraph("Exclude.")
            doc.add_paragraph()
            sent_ = "The service request count is "+str( sr_count)+" which is less than the specified value in GR. Hence excluding."
            doc.add_paragraph(sent_)
            doc.add_page_break()
            excluded_bugs_list.append(var1)
            continue
            
        if ( ex_if_yrs_old ==1 and  (cur_year - int(bug_init_year) >= s6) ):
            p = doc.add_paragraph(" ")
            add_hyperlink(p , var1 , link_str)
            doc.add_paragraph(bug_title)
            doc.add_paragraph("Exclude.")
            doc.add_paragraph()
            sent_ = "The bug is "+str(cur_year - int(bug_init_year) )+" years old. As per GR, exclude bugs if "+str(s6)+" years old."
            doc.add_paragraph(sent_)
            doc.add_page_break()
            excluded_bugs_list.append(var1)
            continue
            
        #if ( ex_if_pst_count==1 and ( sr_count <= s7) ):
        #    p = doc.add_paragraph(" ")
        #    add_hyperlink(p , var1 , link_str)
        #    doc.add_paragraph(bug_title)
        #    doc.add_paragraph("Exclude.")
        #    doc.add_paragraph()
        #    sent_ = "It is a psirt issue with"+str( sr_count)+"service request. As per GR, exclude psirts with"+str(s7)+"or lesser service request."
        #    doc.add_paragraph(sent_)
        #    doc.add_page_break()
        #    excluded_bugs_list.append(var1)
        #    continue     
            
        #if ( ex_pst_yrsold==1 and(cur_year - bug_init_year >= s8)):
        #    p = doc.add_paragraph(" ")
        #    add_hyperlink(p , var1 , link_str)
        #    doc.add_paragraph(bug_title)
        #    doc.add_paragraph("Exclude.")
        #    doc.add_paragraph()
        #    sent_ = "It is a psirt issue with"+str( cur_year - bug_init_year )+ "years old. As per GR, exclude psirts with"+str(s8)+"or more years old."
        #    doc.add_paragraph()
        #    doc.add_page_break()
        #    excluded_bugs_list.append(var1)
        #    continue   
        
        if ( (ex_int_srcount==1) and ( sr_count <= s9) and ( in_fnd_ex_fnd =="Internally" ) ) :
            p = doc.add_paragraph(" ")
            add_hyperlink(p , var1 , link_str)
            doc.add_paragraph(bug_title)
            doc.add_paragraph("Exclude.")
            doc.add_paragraph()
            sent_ = "It is an internally found issue with "+str(sr_count)+ " service requests. As per GR, exclude internally found with "+str(s9)+" or lesser SR."
            doc.add_paragraph(sent_)
            doc.add_page_break()
            excluded_bugs_list.append(var1)
            continue           
            
        #if ( ex_int_yrsold==1 and (cur_year - bug_init_year >= s10)):
        #    p = doc.add_paragraph(" ")
        #    add_hyperlink(p , var1 , link_str)
        #    doc.add_paragraph(bug_title)
        #    doc.add_paragraph("Exclude.")
        #    doc.add_paragraph()
        #    sent_ = "It is an internally found issue with "+ str(cur_year - bug_init_year) +" years old. As per GR, exclude internally found with"+ str(s10)+"or more years old."
        #    doc.add_paragraph(sent_)
        #    doc.add_page_break()
        #    excluded_bugs_list.append(var1)
        #    continue   
            
        #if ( bug_status=="Resolved" and ex_rsldforgivenyrs==1 and (cur_year - bug_init_year >= s11)):
        #    p = doc.add_paragraph(" ")
        #    add_hyperlink(p , var1 , link_str)
        #    doc.add_paragraph(bug_title)
        #    doc.add_paragraph("Exclude.")
        #    doc.add_paragraph()
        #    sent_ = "It is a resolved bug with"+ str(cur_year - bug_init_year)+ "years old. As per GR, exclude resolved bugs with with"+ str(s11) +"or more years old."
        #    doc.add_paragraph(sent_)
        #    doc.add_page_break()
        #    excluded_bugs_list.append(var1)
        #    continue  
        
        if ( ex_intwithSEVimpact==1 and in_fnd_ex_fnd=="Internally" and bug_sev_val>=s12):
            p = doc.add_paragraph(" ")
            add_hyperlink(p , var1 , link_str)
            doc.add_paragraph(bug_title)
            doc.add_paragraph("Exclude.")
            doc.add_paragraph()
            sent_ = "It is an internally found bug with "+ str(bug_sev_val) +" severity level. As per GR, exclude internally found bugs with with "+ str(s12) +" severity level or of lesser impacting levels."
            doc.add_paragraph(sent_)
            doc.add_page_break()
            excluded_bugs_list.append(var1)
            continue   
            
        if ( ex_sev_is==1 and bug_sev_val>=s13):
            p = doc.add_paragraph(" ")
            add_hyperlink(p , var1 , link_str)
            doc.add_paragraph(bug_title)
            doc.add_paragraph("Exclude.")
            doc.add_paragraph()
            sent_ = "It has a severity level of "+ str(bug_sev_val) + ". As per GR, exclude bugs with with " +str(s13)+ " severity level or of lesser impacting levels."
            doc.add_paragraph(sent_)
            doc.add_page_break()
            excluded_bugs_list.append(var1)
            continue         
            
        if ( ex_cl_j_h==1 and (cur_year - int(bug_init_year) >= s14a) and ( sr_count <=s14b) and (bug_status in ['C-Closed','J-Junked','H-Held'] )):
            p = doc.add_paragraph(" ")
            add_hyperlink(p , var1 , link_str)
            doc.add_paragraph(bug_title)
            doc.add_paragraph("Exclude.")
            doc.add_paragraph()
            sent_ = "It is a "+ bug_status + " status bug of "+ str(cur_year - int(bug_init_year) )+ " years old and with "+ str(sr_count)+ " SR. As per GR, excluding."
            doc.add_paragraph(sent_)
            doc.add_page_break()
            excluded_bugs_list.append(var1)
            continue    
        
        
        included_bugs_list.append(var1)
        
        if (only_exc_checker ==1):
            continue
        

        ###
        
        #time.sleep(10)
        mn.click()
        #time.sleep(10)
        text_of_enc = page.locator("//div[@id='enclosures'][@class='col-sm-12 text-lg-left']")
        s = text_of_enc.inner_text() 
        
        topic_list, topic_text = [] , []
        pos = text_of_enc_only_len 
        
        for i in range(topic_count):
            topic_list.append(list_of_enc_only[3+i])
            
        topic_list.append("Top of Page")
        
        for i in range(topic_count):
            next_pos = s.find(topic_list[i+1] , pos)
            topic_text.append(s[pos:next_pos])
            pos=next_pos
            
        p = doc.add_paragraph(" ")
        add_hyperlink(p , var1 , link_str)
        doc.add_paragraph(bug_title)
        p = doc.add_paragraph(" ")
        
        if "Description" in topic_list:
            des_pos=topic_list.index("Description")
            if "static-analysis-" in topic_text[des_pos]:
                st_pos= topic_text[des_pos].index("static-analysis-")
                topic_text[des_pos] = topic_text[des_pos][:st_pos]
            para = doc.add_paragraph(" ")
            para.add_run("Description:").font.highlight_color = WD_COLOR_INDEX.YELLOW
            doc.add_paragraph()
            doc.add_paragraph(topic_text[des_pos+1][11:])
        else:
            para = doc.add_paragraph(" ")
            para.add_run("Description text not available").font.highlight_color = WD_COLOR_INDEX.YELLOW
            doc.add_paragraph()
        doc.add_paragraph()    
        

        if "Release-note" in topic_list:
            rne_pos = topic_list.index("Release-note")
            if "static-analysis-" in topic_text[rne_pos]:
                st_pos= topic_text[rne_pos].index("static-analysis-")
                topic_text[rne_pos] = topic_text[rne_pos][:st_pos]
            strike_pos_rne = topic_text[rne_pos].index("Symptom:")
            para = doc.add_paragraph(" ")
            para.add_run("Release note:").font.highlight_color = WD_COLOR_INDEX.YELLOW
            doc.add_paragraph()
            doc.add_paragraph(topic_text[rne_pos][strike_pos_rne:])
        else:
            para = doc.add_paragraph(" ")
            para.add_run("RNE not available:").font.highlight_color = WD_COLOR_INDEX.YELLOW
            doc.add_paragraph()
        doc.add_paragraph()
        
        
        Email_flag=0
        for i in topic_list:
            if ("Email" in i ) or ("Enail" in i) or("email" in i):
                email_pos = topic_list.index(i)
                if "static-analysis-" in topic_text[email_pos]:
                    st_pos= topic_text[email_pos].index("static-analysis-")
                    topic_text[email_pos] = topic_text[email_pos][:st_pos]
                para = doc.add_paragraph(" ")
                para.add_run("Email found:").font.highlight_color = WD_COLOR_INDEX.YELLOW
                doc.add_paragraph()
                doc.add_paragraph(topic_text[email_pos])
                Email_flag=1
                break
        if Email_flag==0:
            para = doc.add_paragraph(" ")
            para.add_run("Email found:").font.highlight_color = WD_COLOR_INDEX.YELLOW
            doc.add_paragraph()
        doc.add_paragraph()
        
        
        if "Scrubnote" in topic_list:
            scbnt_pos = topic_list.index("Scrubnote")
            if "static-analysis-" in topic_text[scbnt_pos]:
                st_pos= topic_text[scbnt_pos].index("static-analysis-")
                topic_text[scbnt_pos] = topic_text[scbnt_pos][:st_pos]
                para = doc.add_paragraph(" ")
                para.add_run("Scrubnote").font.highlight_color = WD_COLOR_INDEX.YELLOW
                doc.add_paragraph()
            doc.add_paragraph(topic_text[scbnt_pos][9:])
        else:
            para = doc.add_paragraph(" ")
            para.add_run("No Scrubnote found").font.highlight_color = WD_COLOR_INDEX.YELLOW
            doc.add_paragraph()
        doc.add_paragraph()
        
        #print(topic_list)
        #print()
        
        code_review_last_pos =-1
        for i in range(topic_count):
            if "code-review" in topic_list[i]:
                code_review_last_pos = i
                #print("found", topic_list[i],":", topic_text[i][:15])
                
        if code_review_last_pos == -1:
            para= doc.add_paragraph(" ")
            para.add_run("No code-review found").font.highlight_color = WD_COLOR_INDEX.YELLOW
            doc.add_paragraph()
        else:
            para = doc.add_paragraph(" ")
            para.add_run("Code-Review found:").font.highlight_color = WD_COLOR_INDEX.YELLOW
            doc.add_paragraph()
            doc.add_paragraph(topic_text[code_review_last_pos])
            if topic_text[code_review_last_pos] =="":
                doc.add_paragraph("Unable to fetch data from cdets page.")
#         if "static-analysis-" in topic_text[code_review_last_pos]:
#             st_pos= topic_text[code_review_last_pos].index("static-analysis-")
#             topic_text[code_review_last_pos] = topic_text[code_review_last_pos][:st_pos]
#             #strike_pos_code_rev = topic_text[code_review_last_pos].index( ">Do not use \$\$PREFCS unless you\'re 100% sure the issue")
#         doc.add_paragraph(topic_text[code_review_last_pos])
            #print( code_review_last_pos , topic_text[code_review_last_pos])
        doc.add_paragraph()


        if "SS-Eval" in topic_list:
            sseval_pos = topic_list.index("SS-Eval")
            if "static-analysis-" in topic_text[sseval_pos]:
                st_pos= topic_text[sseval_pos].index("static-analysis-")
                topic_text[sseval_pos] = topic_text[sseval_pos][:st_pos]
            para = doc.add_paragraph(" ")
            para.add_run("SS-Eval:").font.highlight_color = WD_COLOR_INDEX.YELLOW
            doc.add_paragraph()
            doc.add_paragraph(topic_text[sseval_pos][7:])
        else:
            para = doc.add_paragraph(" ")
            para.add_run("No SS-Eval found:").font.highlight_color = WD_COLOR_INDEX.YELLOW
            doc.add_paragraph()
        doc.add_paragraph()

        p_tz = doc.add_paragraph(" ")
        p_tz_l = "https://topic.cisco.com/results?queryText="+var1+"&checkedds=news&checkedds=techzone&checkedds=bdb&checkedds=fieldnotice&uncheckedds=cdets&uncheckedds=tstraining&uncheckedds=psirtext&repos=news&repos=techzone&repos=bdb&repos=fieldnotice&isliteralsearch=false&userIdBoost=false&hotWarm=true"        
        add_hyperlink(p_tz , 'TOPIC-SEARCH link' , p_tz_l)
        doc.add_page_break()
        
        
    doc.add_paragraph("Excluded bugs:")
    for e in excluded_bugs_list:
        ee=e+","
        doc.add_paragraph(ee)
    doc.add_paragraph()
    doc.add_paragraph("Included bugs:")
    for c in included_bugs_list:
        ci = c+","
        doc.add_paragraph(ci)
    
#######
#######
#Release-note-diff.txt , Diffs-commit- , FireX_Analysis , 

def open_process():
    doc= Document()
    driver = sync_playwright().start().chromium.launch(headless=False)    
    
    page = driver.new_page()
    bugs_entered_v = bugs_entered.get()
    user_pid = pid.get()
    #print(bugs_entered_v)

    
    ip_bugs = bugs_entered.get()
    
    ip_bugs = ip_bugs.strip()
    ip_bugs = ip_bugs.strip("\n")
    ip_bugs_list = ip_bugs.split("CSC")
    ip_bugs_list = ip_bugs_list[1:]
    bugs_list_len = len(ip_bugs_list)
    for i in range( bugs_list_len ):
        id_ = ip_bugs_list[i]
        id_ = id_.strip("\n")
        id_ = id_.strip(",")
        id_ = id_.strip()
        id_ = "CSC"+id_ 
        ip_bugs_list[i] = id_
        
    only_exc_checker  , ex_if_sr_count , ex_if_yrs_old , ex_int_srcount , ex_intwithSEVimpact , ex_sev_is , ex_cl_j_h  = hi3.get() , hi5.get() , hi6.get() , hi9.get() , hi12.get() , hi13.get() , hi14.get()
    l = [ only_exc_checker , ex_if_sr_count , ex_if_yrs_old , ex_int_srcount , ex_intwithSEVimpact , ex_sev_is , ex_cl_j_h ]
    
    x5,x6,x9,x12,x13,x14a,x14b = s5.get() , s6.get() , s9.get() , s12.get() , s13.get() , s14a.get() , s14b.get()
    
    xl = [int(x5),int(x6),int(x9),int(x12),int(x13),int(x14a),int(x14b)]
    
    #for i in ip_bugs_list:
    start_web_scrapping( ip_bugs_list , doc , page , l , xl , user_pid )
    doc.save("out3.docx")
    #browse_text.set("Finished")



root=tk.Tk()

canvas = tk.Canvas(root,width=850,height = 800)
canvas.grid(columnspan=20,rowspan=17)


#logo 
num1 = random.randint(0, 5)
logo_img_list = [ 'bug_logo.jpg' , 'butterfly_low_poly_color.jpg' , 'techno_gecko.jpg' , 'circuit_butterfly.jpg' , 'binary_butterfly.jpg' , 'box_butterfly.jpg' ]
img_chosen = "Photos-for-bug-clude/"+logo_img_list[num1]
logo = Image.open( img_chosen )
logo = ImageTk.PhotoImage(logo)
logo_label=tk.Label(image=logo)
logo_label.image =logo
logo_label.grid(columnspan =20, column=0, row=0)

t1a = tk.Label(root, text ="Enter",font="Raleway")
t1a.grid( columnspan=1, column=1,row=2)

t1b = tk.Label(root, text =" the",font="Raleway")
t1b.grid( columnspan=1, column=2,row=2)

t1c = tk.Label(root, text =" bug",font="Raleway")
t1c.grid( columnspan=1, column=3,row=2)

t1d = tk.Label(root, text =" ids ",font="Raleway")
t1d.grid( columnspan=1, column=4,row=2)

t1e = tk.Label(root, text ="here :  ",font="Raleway")
t1e.grid( columnspan=1, column=5,row=2)

bugs_entered = StringVar()
bugs_entered_e = Entry(root, textvariable = bugs_entered , width= 100)
bugs_entered_e.grid(columnspan= 13 , column=6 , row=2 , sticky = W) 

t1f = tk.Label(root, text ="  ",font="Raleway")
t1f.grid( sticky=W, columnspan=1, column=20,row=2)


hi3 , hi4,hi5,hi6,hi7,hi8,hi9,hi10,hi11,hi12,hi13,hi14 = IntVar(), IntVar(), IntVar(), IntVar(), IntVar(), IntVar(), IntVar(), IntVar(), IntVar(), IntVar(), IntVar(), IntVar(), 

t3 = tk.Label(root, text ="Show only excluded bugs.",font="Raleway" , anchor="w", justify="left")
t3.grid(sticky = W, columnspan=15, column=1,row=6)

c3 = tk.Checkbutton( variable=hi3, onvalue=1, offvalue=0)
c3.grid( column=0,row=6, columnspan=1)

# t4 = tk.Label(root, text ="Find sentence with “LOC:” in it.",font="Raleway" , anchor="w", justify="left")
# t4.grid(sticky = W, columnspan=13, column= 19 ,row=6)

# c4 = tk.Checkbutton( variable=hi4, onvalue=1, offvalue=0)
# c4.grid( column=18,row=6)

t5 = tk.Label(root, text ="Exclude if SR count is ---- or lesser. ",font="Raleway" , anchor="w", justify="left")
t5.grid(sticky = W, columnspan=15, column=1,row=7)

c5 = tk.Checkbutton( variable=hi5, onvalue=1, offvalue=0)
c5.grid(column=0,row=7, columnspan=1)

s5 = Spinbox(root , from_ = 0 , to_ = 1000 , width=3)
s5.grid( columnspan=1, column=5 , row=7 )


t6 = tk.Label(root, text ="Exclude bugs if ----  or more years old." , font="Raleway" , anchor="w", justify="left")
t6.grid(sticky = W, columnspan=15, column=1 ,row=8)

c6 = tk.Checkbutton( variable=hi6, onvalue=1, offvalue=0)
c6.grid( column= 0 ,row=8 , columnspan = 1)

s6 = Spinbox(root , from_ = 0 , to_ = 1000 , width=3)
s6.grid(    columnspan=1, column= 4 , row=8 )

#t6b = tk.Label(root, text ="or more years old" , font="Raleway" , anchor="w", justify="left")
#t6b.grid(sticky = W, columnspan=10, column= 9 ,row=8)

# #   -  or lesser SR"
#t7 = tk.Label(root, text ="Exclude psirt with"  ,font="Raleway" , anchor="w", justify="left")
#t7.grid(sticky = W, columnspan = 5, column=2,row=8)

#c7 = tk.Checkbutton( variable=hi7, onvalue=1, offvalue=0)
#c7.grid( column=1,row=8)

#s7 = Spinbox(root , from_ = 0 , to_ = 1000 , width=3)
#s7.grid(  columnspan=1, column=7 , row=8 )

#t7b = tk.Label(root, text ="or lesser SR"  ,font="Raleway" , anchor="w", justify="left")
#t7b.grid(sticky = W, columnspan = 5, column=8,row=8)


#t8 = tk.Label(root, text ="Exclude psirt with   ---  or more years old",font="Raleway" , anchor="w", justify="left")
#t8.grid(sticky = W, columnspan=16, column= 19 ,row=8)


#c8 = tk.Checkbutton( variable=hi8, onvalue=1, offvalue=0)
#c8.grid( column= 18,row=8)

#s8 = Spinbox(root , from_ = 0 , to_ = 1000 , width=3)
#s8.grid( columnspan=1, column= 20 , row=8 )


t9 = tk.Label(root, text ="Exclude int found with ---- or lesser SR",font="Raleway" , anchor="w", justify="left")
t9.grid(sticky = W, columnspan=9, column=1,row=9)

c9 = tk.Checkbutton( variable=hi9, onvalue=1, offvalue=0)
c9.grid(column=0,row=9 , columnspan=1)

s9 = Spinbox(root , from_ = 0 , to_ = 1000 , width=3)
s9.grid(  columnspan=1, column=5 , row=9 )

#t10 = tk.Label(root, text ="Exclude int found with   --   or more of years old.",font="Raleway" , anchor="w", justify="left")
#t10.grid(sticky = W, columnspan=13, column= 19  ,row=9)

#c10 = tk.Checkbutton( variable=hi10, onvalue=1, offvalue=0)
#c10.grid(column= 18 ,row=9)

#s10 = Spinbox(root , from_ = 0 , to_ = 1000 , width=3)
#s10.grid( sticky=W, columnspan=1, column= 21 , row=9 )


t20 = tk.Label(root, text ="  ",font="Raleway" , anchor="w", justify="left")
t20.grid( columnspan=10, column=0,row=10)


#t11 = tk.Label(root, text ="Exclude Resolved bugs with   ----  or more years old ",font="Raleway" , anchor="w", justify="left")
#t11.grid(sticky = W, columnspan=18, column=2,row=11)

#c11 = tk.Checkbutton( variable=hi11, onvalue=1, offvalue=0)
#c11.grid(column=1,row=11)

#s11 = Spinbox(root , from_ = 0 , to_ = 1000 , width=3)
#s11.grid( sticky = W, columnspan=1, column=9 , row=11 )


t12 = tk.Label(root, text ="Exclude int found with  ---  or lesser severity impact",font="Raleway" , anchor="w", justify="left")
t12.grid(sticky = W, columnspan=15, column=1,row=11)

c12 = tk.Checkbutton( variable=hi12, onvalue=1, offvalue=0)
c12.grid( column=0,row=11 , columnspan=1)

s12 = Spinbox(root , from_ = 0 , to_ = 1000 , width=3)
s12.grid( columnspan=1, column = 5 , row=11 )

# #

t13 = tk.Label(root, text ="Exclude if severity level is    ---- or of lesser impactng.",font="Raleway" , anchor="w", justify="left")
t13.grid( sticky = W, columnspan=15, column=1,row=12)

c13 = tk.Checkbutton( variable=hi13, onvalue=1, offvalue=0)
c13.grid( column=0,row=12 , columnspan=1)

s13 = Spinbox(root , from_ = 0 , to_ = 1000 , width=3)
s13.grid( sticky = W, columnspan=1, column= 6 , row=12 )

# #

t14 = tk.Label(root, text ="Exclude Closed/ Junked /Held bugs with  ....  or more years old with    ..   or lesser SR",font="Raleway" , anchor="w", justify="left")
t14.grid(sticky = W, columnspan=9, column=1,row=13 )

c14 = tk.Checkbutton( variable=hi14, onvalue=1, offvalue=0)
c14.grid( columnspan=1 , column=0 , row=13)

s14a = Spinbox(root , from_ = 0 , to_ = 1000 , width=3)
s14a.grid( sticky = E , columnspan=1, column=6 , row=13 )

s14b = Spinbox(root , from_ = 0 , to_ = 1000 , width=3)
s14b.grid( sticky =W ,  columnspan=1, column=9 , row=13)

t15 = tk.Label(root, text ="Enter your your PID id: ",font="Raleway")
t15.grid( columnspan=5, column=1,row=14)

pid = StringVar()
pid_e = Entry(root, textvariable = bugs_entered , width= 30)
pid_e.grid(columnspan= 4 , column=6 , row=14 , sticky = W)



browse_text = tk.StringVar()
browse_text.set("Submit")
browse_btn = tk.Button(root, textvariable=browse_text, command=lambda:open_process(),font ="Raleway", bg="green", fg="white",height=2,width=12 )
browse_btn.grid(columnspan=10,column=1,row=16)

#GUI CODE FINISHED

root.mainloop()